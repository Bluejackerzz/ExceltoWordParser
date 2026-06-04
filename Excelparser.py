import os
import re
import sys
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# GUI Imports
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import tkinter.scrolledtext as scrolledtext

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel to Word Parser V.1.2")
        self.root.geometry("850x750")
        self.root.minsize(800, 650)
        
        # State variables
        self.excel_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.date_format = tk.StringVar(value="DD-MM-YYYY")
        self.processing_mode = tk.StringVar(value="all") # "all", "single", "range"
        self.single_row_num = tk.StringVar(value="1")
        self.range_start = tk.StringVar(value="1")
        self.range_end = tk.StringVar(value="5")
        self.filename_column = tk.StringVar()
        self.font_name = tk.StringVar(value="Tahoma")
        self.font_size = tk.StringVar(value="10")
        
        # Available Excel columns (populated after loading Excel)
        self.excel_columns = []
        
        # Dynamic mapping lists
        self.mappings = []       # Word to Excel mappings
        self.custom_fills = []   # Word to Custom Static Constant fills
        
        self.build_ui()
        self.load_default_configuration()

    def log(self, message):
        """Appends status messages to the on-screen terminal box."""
        self.log_widget.configure(state='normal')
        self.log_widget.insert(tk.END, message + "\n")
        self.log_widget.see(tk.END)
        self.log_widget.configure(state='disabled')
        self.root.update_idletasks()

    def build_ui(self):
        # Create Main Frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # --- File Selection Panel ---
        files_lf = ttk.LabelFrame(main_frame, text=" 1. File Selection ", padding="10")
        files_lf.pack(fill=tk.X, pady=5)
        
        # Excel File
        ttk.Label(files_lf, text="Excel File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(files_lf, textvariable=self.excel_path, width=70).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(files_lf, text="Browse...", command=self.browse_excel).grid(row=0, column=2, padx=5, pady=5)
        
        # Word Template File
        ttk.Label(files_lf, text="Word Template:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(files_lf, textvariable=self.template_path, width=70).grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(files_lf, text="Browse...", command=self.browse_template).grid(row=1, column=2, padx=5, pady=5)
        
        # Output Folder
        ttk.Label(files_lf, text="Output Directory:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(files_lf, textvariable=self.output_dir, width=70).grid(row=2, column=1, padx=5, pady=5)
        ttk.Button(files_lf, text="Browse...", command=self.browse_output).grid(row=2, column=2, padx=5, pady=5)

        # --- Configurations and Mappings Panel ---
        middle_paned = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        middle_paned.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Left side configuration frame
        config_frame = ttk.Frame(middle_paned, padding="5")
        middle_paned.add(config_frame, weight=1)
        
        # Right side Notebook (Tabs) for Mappings
        right_notebook_frame = ttk.Frame(middle_paned, padding="5")
        middle_paned.add(right_notebook_frame, weight=1)
        
        self.mappings_notebook = ttk.Notebook(right_notebook_frame)
        self.mappings_notebook.pack(fill=tk.BOTH, expand=True)
        
        # Tab 1: Excel Mapping Panel
        self.tab_excel = ttk.Frame(self.mappings_notebook, padding="10")
        self.mappings_notebook.add(self.tab_excel, text=" Excel Mappings ")
        
        # Tab 2: Custom Constant Fill Panel
        self.tab_custom = ttk.Frame(self.mappings_notebook, padding="10")
        self.mappings_notebook.add(self.tab_custom, text=" Custom Constant Fills ")
        
        # --- Config elements (Left Panel) ---
        options_lf = ttk.LabelFrame(config_frame, text=" Options ", padding="10")
        options_lf.pack(fill=tk.BOTH, expand=True)
        
        # Date Format Dropdown
        ttk.Label(options_lf, text="Date Format:").grid(row=0, column=0, sticky=tk.W, pady=5)
        date_combo = ttk.Combobox(options_lf, textvariable=self.date_format, width=15, state="readonly")
        date_combo['values'] = ("DD-MM-YYYY", "MM-DD-YYYY", "YYYY-MM-DD", "DD/MM/YYYY", "MM/DD/YYYY")
        date_combo.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Row selection mode
        ttk.Label(options_lf, text="Processing Rows:").grid(row=1, column=0, sticky=tk.W, pady=5)
        
        mode_frame = ttk.Frame(options_lf)
        mode_frame.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Radiobutton(mode_frame, text="All Rows", variable=self.processing_mode, value="all", command=self.toggle_row_inputs).pack(anchor=tk.W)
        ttk.Radiobutton(mode_frame, text="Single Row", variable=self.processing_mode, value="single", command=self.toggle_row_inputs).pack(anchor=tk.W)
        
        self.single_entry_frame = ttk.Frame(mode_frame)
        self.single_entry_frame.pack(anchor=tk.W, padx=15)
        ttk.Label(self.single_entry_frame, text="Row Index (1-based):").pack(side=tk.LEFT)
        self.single_row_spin = ttk.Spinbox(self.single_entry_frame, from_=1, to=10000, textvariable=self.single_row_num, width=6)
        self.single_row_spin.pack(side=tk.LEFT, padx=5)
        
        ttk.Radiobutton(mode_frame, text="Row Range", variable=self.processing_mode, value="range", command=self.toggle_row_inputs).pack(anchor=tk.W)
        
        self.range_entry_frame = ttk.Frame(mode_frame)
        self.range_entry_frame.pack(anchor=tk.W, padx=15)
        ttk.Label(self.range_entry_frame, text="From:").pack(side=tk.LEFT)
        self.range_start_spin = ttk.Spinbox(self.range_entry_frame, from_=1, to=10000, textvariable=self.range_start, width=5)
        self.range_start_spin.pack(side=tk.LEFT, padx=5)
        ttk.Label(self.range_entry_frame, text="To:").pack(side=tk.LEFT)
        self.range_end_spin = ttk.Spinbox(self.range_entry_frame, from_=1, to=10000, textvariable=self.range_end, width=5)
        self.range_end_spin.pack(side=tk.LEFT, padx=5)
        
        self.toggle_row_inputs() # Initialize state
        
        # File Naming Selector
        ttk.Label(options_lf, text="File Name Column:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.filename_combo = ttk.Combobox(options_lf, textvariable=self.filename_column, width=20, state="readonly")
        self.filename_combo.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)

        # Font Family Row
        ttk.Label(options_lf, text="Font Family:").grid(row=3, column=0, sticky=tk.W, pady=5)
        font_combo = ttk.Combobox(options_lf, textvariable=self.font_name, width=20)
        font_combo['values'] = ("Tahoma", "Arial", "Calibri", "Times New Roman", "Segoe UI", "Arial Black", "Courier New", "Georgia", "Verdana")
        font_combo.grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)

        # Font Size Row
        ttk.Label(options_lf, text="Font Size:").grid(row=4, column=0, sticky=tk.W, pady=5)
        size_spin = ttk.Spinbox(options_lf, from_=6, to=72, textvariable=self.font_size, width=6)
        size_spin.grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)
        
        # --- TAB 1: Excel Mapping Elements ---
        control_mapping_frame = ttk.Frame(self.tab_excel)
        control_mapping_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(control_mapping_frame, text="Add Field Mapping", command=self.add_mapping_row).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_mapping_frame, text="Clear All", command=self.clear_mappings).pack(side=tk.LEFT, padx=5)
        
        # Headers for Excel Mapping
        header_frame_excel = ttk.Frame(self.tab_excel)
        header_frame_excel.pack(fill=tk.X, pady=2)
        ttk.Label(header_frame_excel, text="Word Template Field Label", font="-weight bold", width=25).pack(side=tk.LEFT, padx=5)
        ttk.Label(header_frame_excel, text="Excel Column", font="-weight bold", width=25).pack(side=tk.LEFT, padx=5)
        
        # Scrollable container for Excel Mappings
        self.canvas_excel = tk.Canvas(self.tab_excel, borderwidth=0, highlightthickness=0)
        self.scroll_frame_excel = ttk.Frame(self.canvas_excel)
        self.scrollbar_excel = ttk.Scrollbar(self.tab_excel, orient="vertical", command=self.canvas_excel.yview)
        self.canvas_excel.configure(yscrollcommand=self.scrollbar_excel.set)
        
        self.scrollbar_excel.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas_excel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.canvas_excel.create_window((0,0), window=self.scroll_frame_excel, anchor="nw", tags="self.scroll_frame_excel")
        self.scroll_frame_excel.bind("<Configure>", lambda event, canvas=self.canvas_excel: self.on_frame_configure(canvas))
        
        # --- TAB 2: Custom Constant Fill Elements ---
        control_custom_frame = ttk.Frame(self.tab_custom)
        control_custom_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(control_custom_frame, text="Add Custom Fill", command=self.add_custom_fill_row).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_custom_frame, text="Clear All", command=self.clear_custom_fills).pack(side=tk.LEFT, padx=5)
        
        # Headers for Custom Fills
        header_frame_custom = ttk.Frame(self.tab_custom)
        header_frame_custom.pack(fill=tk.X, pady=2)
        ttk.Label(header_frame_custom, text="Word Template Field Label", font="-weight bold", width=25).pack(side=tk.LEFT, padx=5)
        ttk.Label(header_frame_custom, text="Custom Text Value", font="-weight bold", width=25).pack(side=tk.LEFT, padx=5)
        
        # Scrollable container for Custom Fills
        self.canvas_custom = tk.Canvas(self.tab_custom, borderwidth=0, highlightthickness=0)
        self.scroll_frame_custom = ttk.Frame(self.canvas_custom)
        self.scrollbar_custom = ttk.Scrollbar(self.tab_custom, orient="vertical", command=self.canvas_custom.yview)
        self.canvas_custom.configure(yscrollcommand=self.scrollbar_custom.set)
        
        self.scrollbar_custom.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas_custom.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.canvas_custom.create_window((0,0), window=self.scroll_frame_custom, anchor="nw", tags="self.scroll_frame_custom")
        self.scroll_frame_custom.bind("<Configure>", lambda event, canvas=self.canvas_custom: self.on_frame_configure(canvas))

        # --- Process and Terminal Log Output Panel ---
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=5)
        
        self.process_btn = ttk.Button(action_frame, text="START PROCESSING", command=self.start_processing)
        self.process_btn.pack(fill=tk.X, ipady=5)
        
        log_lf = ttk.LabelFrame(main_frame, text=" Run Status Logs ", padding="5")
        log_lf.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.log_widget = scrolledtext.ScrolledText(log_lf, height=8, state='disabled', wrap=tk.WORD, font=("Consolas", 9))
        self.log_widget.pack(fill=tk.BOTH, expand=True)

    def on_frame_configure(self, canvas):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def toggle_row_inputs(self):
        """Enables or disables row bounds inputs depending on selected option."""
        mode = self.processing_mode.get()
        if mode == "all":
            self.set_widget_state(self.single_entry_frame, "disabled")
            self.set_widget_state(self.range_entry_frame, "disabled")
        elif mode == "single":
            self.set_widget_state(self.single_entry_frame, "normal")
            self.set_widget_state(self.range_entry_frame, "disabled")
        elif mode == "range":
            self.set_widget_state(self.single_entry_frame, "disabled")
            self.set_widget_state(self.range_entry_frame, "normal")

    def set_widget_state(self, parent, state):
        """Helper to recursively set widget states within a container frame."""
        for child in parent.winfo_children():
            try:
                child.configure(state=state)
            except Exception:
                pass

    def browse_excel(self):
        filename = filedialog.askopenfilename(
            title="Select Excel Data File",
            filetypes=[("Excel Files", "*.xlsx *.xls"), ("All Files", "*.*")]
        )
        if filename:
            self.excel_path.set(filename)
            self.log(f"Selected Excel file: {filename}")
            self.load_excel_headers(filename)

    def browse_template(self):
        filename = filedialog.askopenfilename(
            title="Select Microsoft Word Template File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.template_path.set(filename)
            self.log(f"Selected Word Template file: {filename}")

    def browse_output(self):
        directory = filedialog.askdirectory(title="Select Output Directory")
        if directory:
            self.output_dir.set(directory)
            self.log(f"Selected Output Directory: {directory}")

    def load_excel_headers(self, filepath):
        """Attempts to parse headers from the Excel file and updates dynamic dropdown selections."""
        try:
            df = pd.read_excel(filepath, nrows=2)
            self.excel_columns = list(df.columns)
            self.log(f"Loaded Excel columns: {', '.join(self.excel_columns)}")
            
            # Update file naming selector options
            self.filename_combo['values'] = self.excel_columns
            if "Project name" in self.excel_columns:
                self.filename_column.set("Project name")
            elif len(self.excel_columns) > 0:
                self.filename_column.set(self.excel_columns[0])
                
            # Update all active dynamic dropdown structures
            for m in self.mappings:
                m["excel_combo"]['values'] = self.excel_columns
        except Exception as e:
            self.log(f"Error reading Excel columns: {e}")
            messagebox.showerror("Excel Error", f"Unable to read file headers:\n{e}")

    def load_default_configuration(self):
        """Pre-populates the paths and mappings to establish default configurations on launch."""
        # Check if workspace template files exist, and pre-fill them
        default_template = "test.docx"
        default_excel = "test.xlsx"
        default_out = "Filled_Reports"

        if os.path.exists(default_template):
            self.template_path.set(os.path.abspath(default_template))
        if os.path.exists(default_excel):
            self.excel_path.set(os.path.abspath(default_excel))
            self.load_excel_headers(os.path.abspath(default_excel))
        
        self.output_dir.set(os.path.abspath(default_out))
        
        # Configure initial target fields (Excel Mappings)
        defaults_excel = [
            ("Project Name", "Project name"),
            ("Service Contract No.", "Contract No."),
            ("Customer Name", "End User"),
            ("Service Start date", "Service start date"),
            ("Service End date", "Service end date")
        ]
        for word, excel in defaults_excel:
            self.add_mapping_row(word, excel)
            
        # Configure initial custom fills (Constant Values)
        defaults_custom = [
            ("H3C Service Manager", "Mr John Doe"),
            ("H3C Project Manager", "Mrs Jane Smith"),
            ("Delivery Engineer", "Alex Tech"),
            ("Service Site", "Jakarta Headquarters")
        ]
        for word, const_val in defaults_custom:
            self.add_custom_fill_row(word, const_val)

    def add_mapping_row(self, initial_word="", initial_excel=""):
        """Creates a dynamic grid row on the Excel mapping panel frame."""
        row_frame = ttk.Frame(self.scroll_frame_excel)
        row_frame.pack(fill=tk.X, pady=2, expand=True)
        
        word_var = tk.StringVar(value=initial_word)
        excel_var = tk.StringVar(value=initial_excel)
        
        word_entry = ttk.Entry(row_frame, textvariable=word_var, width=25)
        word_entry.pack(side=tk.LEFT, padx=5)
        
        excel_combo = ttk.Combobox(row_frame, textvariable=excel_var, values=self.excel_columns, width=25)
        excel_combo.pack(side=tk.LEFT, padx=5)
        
        del_btn = ttk.Button(row_frame, text="X", width=3, command=lambda f=row_frame: self.remove_mapping_row(f))
        del_btn.pack(side=tk.LEFT, padx=5)
        
        self.mappings.append({
            "word_var": word_var,
            "excel_var": excel_var,
            "frame": row_frame,
            "excel_combo": excel_combo
        })

    def remove_mapping_row(self, frame_to_remove):
        """Finds and destroys a dynamic mapping frame reference."""
        for m in list(self.mappings):
            if m["frame"] == frame_to_remove:
                m["frame"].destroy()
                self.mappings.remove(m)
                break

    def clear_mappings(self):
        for m in list(self.mappings):
            m["frame"].destroy()
        self.mappings.clear()

    def add_custom_fill_row(self, initial_word="", initial_text=""):
        """Creates a dynamic grid row on the Custom Fills panel frame."""
        row_frame = ttk.Frame(self.scroll_frame_custom)
        row_frame.pack(fill=tk.X, pady=2, expand=True)
        
        word_var = tk.StringVar(value=initial_word)
        text_var = tk.StringVar(value=initial_text)
        
        word_entry = ttk.Entry(row_frame, textvariable=word_var, width=25)
        word_entry.pack(side=tk.LEFT, padx=5)
        
        text_entry = ttk.Entry(row_frame, textvariable=text_var, width=25)
        text_entry.pack(side=tk.LEFT, padx=5)
        
        del_btn = ttk.Button(row_frame, text="X", width=3, command=lambda f=row_frame: self.remove_custom_fill_row(f))
        del_btn.pack(side=tk.LEFT, padx=5)
        
        self.custom_fills.append({
            "word_var": word_var,
            "text_var": text_var,
            "frame": row_frame
        })

    def remove_custom_fill_row(self, frame_to_remove):
        """Finds and destroys a dynamic custom fill frame reference."""
        for cf in list(self.custom_fills):
            if cf["frame"] == frame_to_remove:
                cf["frame"].destroy()
                self.custom_fills.remove(cf)
                break

    def clear_custom_fills(self):
        for cf in list(self.custom_fills):
            cf["frame"].destroy()
        self.custom_fills.clear()

    def start_processing(self):
        """Initiates parser data structures and handles runtime validation execution."""
        excel_fp = self.excel_path.get()
        template_fp = self.template_path.get()
        out_dir = self.output_dir.get()
        
        if not excel_fp or not os.path.exists(excel_fp):
            messagebox.showerror("Error", "Please select a valid Excel file path.")
            return
        if not template_fp or not os.path.exists(template_fp):
            messagebox.showerror("Error", "Please select a valid Word template file path.")
            return
        if not out_dir:
            messagebox.showerror("Error", "Please specify a directory to save output documents.")
            return
            
        if not os.path.exists(out_dir):
            try:
                os.makedirs(out_dir)
            except Exception as e:
                messagebox.showerror("Error", f"Could not create output directory:\n{e}")
                return

        # Prepare Excel-to-Word mapping configurations
        field_mapping = {}
        for m in self.mappings:
            word_lbl = m["word_var"].get().strip()
            excel_col = m["excel_var"].get().strip()
            if word_lbl and excel_col:
                norm_lbl = re.sub(r'[^a-zA-Z0-9]', '', word_lbl).lower()
                field_mapping[norm_lbl] = {
                    "word_original": word_lbl,
                    "excel_original": excel_col
                }

        # Prepare Custom Constant mapping configurations
        custom_mapping = {}
        for cf in self.custom_fills:
            word_lbl = cf["word_var"].get().strip()
            const_val = cf["text_var"].get().strip()
            if word_lbl:
                norm_lbl = re.sub(r'[^a-zA-Z0-9]', '', word_lbl).lower()
                custom_mapping[norm_lbl] = const_val

        if not field_mapping and not custom_mapping:
            messagebox.showerror("Error", "Please configure at least one Excel mapping or Custom Constant Fill.")
            return

        self.process_btn.configure(state='disabled')
        self.log("=== Started Mail-Merge Execution ===")
        
        try:
            df = pd.read_excel(excel_fp)
        except Exception as e:
            self.log(f"Error: Unable to load Excel spreadsheet data: {e}")
            messagebox.showerror("Error", f"Failed to load spreadsheet: {e}")
            self.process_btn.configure(state='normal')
            return

        total_rows = len(df)
        self.log(f"Spreadsheet loaded. Found {total_rows} total rows of data.")
        
        # Parse processing indices boundaries
        mode = self.processing_mode.get()
        start_idx = 0
        end_idx = total_rows
        
        if mode == "single":
            try:
                single_idx = int(self.single_row_num.get()) - 1
                if single_idx < 0 or single_idx >= total_rows:
                    raise ValueError(f"Value must be between 1 and {total_rows}")
                start_idx = single_idx
                end_idx = single_idx + 1
            except ValueError as e:
                messagebox.showerror("Invalid Row Index", f"Verify row is within bounds: {e}")
                self.process_btn.configure(state='normal')
                return
        elif mode == "range":
            try:
                start_num = int(self.range_start.get()) - 1
                end_num = int(self.range_end.get())
                if start_num < 0 or end_num > total_rows or start_num >= end_num:
                    raise ValueError("Verify 'From' is less than 'To' and remains within bounds.")
                start_idx = start_num
                end_idx = end_num
            except ValueError as e:
                messagebox.showerror("Invalid Range Range", f"Adjust boundary options:\n{e}")
                self.process_btn.configure(state='normal')
                return

        date_fmt = self.date_format.get()
        naming_col = self.filename_column.get()
        
        # Extract selected Font Properties from dynamic configurations
        font_name_val = self.font_name.get().strip() or "Tahoma"
        try:
            font_size_val = float(self.font_size.get())
        except ValueError:
            font_size_val = 10.0

        self.log(f"Processing row index bounds: {start_idx + 1} to {end_idx}")
        self.log(f"Applied Font Formatting: {font_name_val}, Size {font_size_val} pt")
        processed_count = 0
        
        for i in range(start_idx, end_idx):
            row = df.iloc[i]
            
            # File generation names
            base_name = str(row.get(naming_col, f"Doc_Row_{i + 1}")).strip()
            if not base_name or pd.isna(row.get(naming_col)):
                base_name = f"Doc_Row_{i + 1}"
            
            # Clean filename strings for security boundaries
            safe_base_name = re.sub(r'[\\/*?:"<>|]', "_", base_name)
            output_filename = f"{safe_base_name}.docx"
            output_filepath = os.path.join(out_dir, output_filename)
            
            # Prevent overwriting logic by appending increment indexes if the file exists
            if os.path.exists(output_filepath):
                count_idx = 1
                while os.path.exists(os.path.join(out_dir, f"{safe_base_name}_{count_idx}.docx")):
                    count_idx += 1
                output_filename = f"{safe_base_name}_{count_idx}.docx"
                output_filepath = os.path.join(out_dir, output_filename)
                
            self.log(f"Processing row {i + 1}/{total_rows} -> Outputting to: {output_filename}")
            
            try:
                self.fill_document_template(template_fp, row, field_mapping, custom_mapping, date_fmt, output_filepath, font_name_val, font_size_val)
                processed_count += 1
            except Exception as e:
                self.log(f"Critical execution error on row {i + 1}: {e}")
                
        self.log(f"=== Execution Terminated. Success Count: {processed_count} ===")
        messagebox.showinfo("Success", f"Mail Merge Completed Successfully.\nGenerated {processed_count} Word document files.")
        self.process_btn.configure(state='normal')

    def fill_document_template(self, doc_path, row_data, field_mapping, custom_mapping, date_fmt, output_path, font_name, font_size):
        """Processes and merges Word document formatting layout configurations."""
        doc = Document(doc_path)
        
        # Format mapping structures to extract correct values and set dates accordingly
        formatted_row_values = {}
        for norm_lbl, m_info in field_mapping.items():
            excel_col = m_info["excel_original"]
            raw_val = row_data.get(excel_col, "")
            
            # Use rules to format dates cleanly
            is_date_field = "date" in norm_lbl
            if is_date_field:
                if pd.isna(raw_val) or raw_val is None or str(raw_val).strip() == "":
                    val_str = ""
                else:
                    try:
                        dt = pd.to_datetime(raw_val)
                        # Map selections
                        format_str = "%d-%m-%Y"
                        if date_fmt == "MM-DD-YYYY":
                            format_str = "%m-%d-%Y"
                        elif date_fmt == "YYYY-MM-DD":
                            format_str = "%Y-%m-%d"
                        elif date_fmt == "DD/MM/YYYY":
                            format_str = "%d/%m/%y"
                        elif date_fmt == "MM/DD/YYYY":
                            format_str = "%m/%d/%y"
                        
                        val_str = dt.strftime(format_str)
                    except Exception:
                        val_str = str(raw_val).strip()
            else:
                val_str = str(raw_val).strip() if pd.notna(raw_val) else ""
            
            formatted_row_values[norm_lbl] = val_str

        # Add custom constant fills (these do not rely on dynamic Excel column values)
        for norm_lbl, const_val in custom_mapping.items():
            formatted_row_values[norm_lbl] = const_val

        # 1. Process all tables in the document
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate physical cells in case of merged structures
                unique_cells = []
                for cell in row.cells:
                    if cell not in unique_cells:
                        unique_cells.append(cell)
                
                num_unique = len(unique_cells)
                for cell_idx, cell in enumerate(unique_cells):
                    cell_text = cell.text.strip()
                    normalized_cell = re.sub(r'[^a-zA-Z0-9]', '', cell_text).lower()
                    
                    # Target correct fields
                    for norm_lbl, val_str in formatted_row_values.items():
                        if normalized_cell == norm_lbl:
                            if cell_idx + 1 < num_unique:
                                target_cell = unique_cells[cell_idx + 1]
                                
                                # Process cell paragraphs to write values
                                if len(target_cell.paragraphs) > 0:
                                    p = target_cell.paragraphs[0]
                                    p.text = ""  # Erases initial text
                                    run = p.add_run(val_str)
                                    apply_font_formatting(run, font_name, font_size)
                                    
                                    # Erase remaining paragraphs inside the cell bounds
                                    for extra_p in target_cell.paragraphs[1:]:
                                        p_element = extra_p._element
                                        p_element.getparent().remove(p_element)
                                else:
                                    p = target_cell.add_paragraph()
                                    run = p.add_run(val_str)
                                    apply_font_formatting(run, font_name, font_size)
                                break

        # 2. Process paragraphs in main body (if containing templates in double brackets, e.g. {{Project name}})
        for paragraph in doc.paragraphs:
            for norm_lbl, m_info in field_mapping.items():
                excel_col = m_info["excel_original"]
                placeholder = f"{{{{{excel_col}}}}}"
                if placeholder in paragraph.text:
                    val_str = formatted_row_values[norm_lbl]
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, val_str)
                            apply_font_formatting(run, font_name, font_size)

        doc.save(output_path)


def apply_font_formatting(run, font_name, font_size):
    """
    Applies the specified font name and size in points to the specified text run.
    Uses XML elements to override Word's language-specific font fallbacks.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    
    # Access structural run properties to assign font definitions explicitly
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    
    # Configure properties across all encoding sets to maintain consistency in Word
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def main():
    root = tk.Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()